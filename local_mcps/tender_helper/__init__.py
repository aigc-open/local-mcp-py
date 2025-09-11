from mcp.server.fastmcp import FastMCP
import os
import json
from typing import List, Dict, Any


_tender_outline_generator_prompt = """
你是一个标书编写专家，请根据给出的招标需求，生成一个详细的标书大纲

## 标书大纲生成要求：
- 你精通各种类型的标书编写：工程类、服务类、货物类、技术类等
- 你了解标书的结构规范和评分要点
- 请根据招标文件的具体要求，生成针对性的标书大纲
- 大纲应该结构清晰、层次分明、符合招标要求
- 每个章节都要有明确的目标和内容说明

## 支持的标书类型：
- **工程类标书**：建筑工程、市政工程、装修工程等
- **服务类标书**：咨询服务、运维服务、培训服务等
- **货物类标书**：设备采购、软件采购、材料采购等
- **技术类标书**：软件开发、系统集成、技术方案等
- **综合类标书**：工程+服务、技术+货物等复合类型

## 标书大纲结构要求：
1. **投标函及投标函附录** - 投标承诺和基本信息
2. **法定代表人身份证明** - 企业资质证明
3. **授权委托书** - 投标授权文件
4. **投标保证金** - 保证金缴纳证明
5. **技术标** - 技术方案和实施计划
6. **商务标** - 商务条款和报价
7. **资格审查资料** - 企业资质和业绩
8. **其他材料** - 招标文件要求的其他内容

## 输出格式要求：
- 使用JSON格式输出大纲结构
- 每个章节包含：章节编号、章节标题、内容要点、页数估算
- 提供章节间的逻辑关系和依赖关系
- 标注重点章节和评分权重

## 大纲JSON结构示例：
```json
{
  "tender_info": {
    "project_name": "项目名称",
    "tender_type": "标书类型",
    "total_pages": "预估总页数"
  },
  "outline": [
    {
      "chapter_id": "01",
      "chapter_title": "投标函及投标函附录",
      "content_points": ["投标承诺", "项目理解", "服务承诺"],
      "estimated_pages": 5,
      "weight": "必备项",
      "dependencies": []
    }
  ]
}
```

请告诉我：
1. 招标项目名称和类型
2. 主要技术要求或服务内容
3. 特殊要求或评分重点
4. 预算范围和工期要求

我将为你生成一个专业、完整、符合要求的标书大纲。
"""

_tender_content_generator_prompt = """
你是一个标书编写专家，请根据提供的标书大纲和具体章节要求，生成详细的标书内容

## 内容生成要求：
- 你精通标书编写的各种规范和技巧
- 你能够根据招标要求生成针对性的投标内容
- 请根据大纲的章节要求，生成详细、专业、有说服力的内容
- 内容应该突出投标方的优势和能力
- 需要考虑评分标准和竞争策略

## 内容生成指导：
1. **技术方案**：详细的技术实现方案和创新点
2. **实施计划**：科学合理的项目实施计划和里程碑
3. **质量保证**：完善的质量管理体系和保证措施
4. **团队配置**：专业的项目团队和人员配置
5. **风险控制**：全面的风险识别和应对措施
6. **售后服务**：完善的售后服务体系和承诺

## 支持的内容类型：
- **技术标内容**：技术方案、实施方案、质量保证等
- **商务标内容**：报价明细、商务条款、付款方式等
- **资格证明**：企业资质、项目业绩、团队介绍等
- **投标文件**：投标函、承诺书、授权书等
- **附件材料**：证书复印件、业绩证明、案例介绍等

## 输出要求：
- 使用Markdown格式
- 包含必要的表格、图表、流程图的描述
- 提供具体的数据和案例支撑
- 突出竞争优势和差异化特点
- 符合招标文件的格式要求

请提供：
1. 标书大纲结构
2. 需要生成内容的具体章节
3. 该章节的重点要求和评分标准
4. 投标方的基本信息和优势
5. 任何特殊要求或约束条件

我将为你生成专业、有说服力、符合要求的标书内容。
"""

_tender_outline_adjuster_prompt = """
你是一个标书编写专家，请根据用户的反馈和要求，调整和优化标书大纲

## 大纲调整要求：
- 你能够理解用户对大纲的修改需求
- 你可以增加、删除、修改、重排大纲章节
- 调整后的大纲要保持逻辑性和完整性
- 确保调整后的大纲符合标书规范

## 支持的调整操作：
1. **增加章节**：在指定位置添加新的章节
2. **删除章节**：移除不需要的章节
3. **修改章节**：更改章节标题、内容要点、页数等
4. **重排章节**：调整章节顺序和层级关系
5. **合并章节**：将多个章节合并为一个
6. **拆分章节**：将一个章节拆分为多个子章节

## 调整指导原则：
- 保持标书的完整性和规范性
- 确保重要章节不遗漏
- 合理分配各章节的篇幅
- 突出项目的重点和优势
- 符合招标文件的要求

## 输入格式：
请提供：
1. 当前的标书大纲（JSON格式）
2. 具体的调整要求和说明
3. 调整的原因和目标

## 输出格式：
- 返回调整后的完整大纲（JSON格式）
- 说明主要的调整内容和理由
- 提供调整建议和注意事项

我将帮你优化标书大纲，确保其更好地满足招标要求和竞争需要。
"""

_tender_file_manager_prompt = """
你是一个标书文件管理专家，请根据标书大纲管理各个章节文件

## 文件管理功能：
- 根据大纲创建章节文件结构
- 管理各章节的Markdown文件
- 跟踪文件的生成状态和更新时间
- 提供文件合并和导出功能

## 文件结构规范：
```
tender_project/
├── outline.json          # 标书大纲文件
├── chapters/             # 章节内容目录
│   ├── 01_投标函.md
│   ├── 02_技术方案.md
│   ├── 03_实施计划.md
│   └── ...
├── attachments/          # 附件目录
│   ├── certificates/     # 证书文件
│   ├── cases/           # 案例材料
│   └── images/          # 图片文件
└── output/              # 输出目录
    ├── tender_full.md   # 合并后的完整文档
    └── tender_full.docx # Word格式文档
```

## 支持的操作：
1. **创建项目结构**：根据大纲创建文件夹和文件
2. **生成章节文件**：为每个章节创建对应的Markdown文件
3. **更新文件状态**：跟踪各文件的编写进度
4. **合并章节内容**：将所有章节合并为完整文档
5. **导出Word文档**：转换为标准的Word格式

## 文件状态管理：
- **未开始**：章节文件已创建但未编写内容
- **进行中**：章节内容正在编写
- **已完成**：章节内容编写完成
- **需修改**：章节内容需要修改完善

请告诉我需要执行的文件管理操作，我将帮你高效管理标书文件。
"""

_tender_word_exporter_prompt = """
你是一个文档格式转换专家，请将标书Markdown内容转换为专业的Word文档

## Word导出要求：
- 转换Markdown格式为Word格式
- 保持文档的格式和样式
- 添加专业的标书样式模板
- 生成目录和页码
- 确保表格和图片正确显示

## Word样式规范：
1. **页面设置**：A4纸张，2.5cm边距
2. **字体规范**：正文宋体12号，标题黑体
3. **段落格式**：1.5倍行距，首行缩进2字符
4. **标题样式**：多级标题自动编号
5. **表格样式**：统一的表格边框和对齐
6. **页眉页脚**：项目名称和页码

## 导出功能：
- 自动生成封面页
- 插入目录页
- 设置页眉页脚
- 应用标书样式模板
- 处理图片和表格
- 生成最终的docx文件

## 使用方法：
1. 提供合并后的完整Markdown内容
2. 指定项目名称和基本信息
3. 选择标书样式模板
4. 生成专业的Word文档

我将帮你生成符合标书规范的专业Word文档。
"""

def register_tender_helper_tools(mcp: FastMCP):
    """注册标书生成助手工具"""
    
    @mcp.tool()
    def tender_outline_generator() -> str:
        """生成标书大纲 - 根据招标需求生成详细的标书大纲结构"""
        return _tender_outline_generator_prompt.strip()
    
    @mcp.tool()
    def tender_outline_adjuster() -> str:
        """调整标书大纲 - 根据用户反馈调整和优化标书大纲"""
        return _tender_outline_adjuster_prompt.strip()
    
    @mcp.tool()
    def tender_content_generator() -> str:
        """生成标书内容 - 根据大纲生成具体章节的详细内容"""
        return _tender_content_generator_prompt.strip()
    
    @mcp.tool()
    def tender_file_manager() -> str:
        """管理标书文件 - 创建文件结构、管理章节文件、合并内容"""
        return _tender_file_manager_prompt.strip()
    
    @mcp.tool()
    def tender_word_exporter() -> str:
        """导出Word文档 - 将标书内容转换为专业的Word格式文档"""
        return _tender_word_exporter_prompt.strip()
    
    @mcp.tool()
    def create_tender_project_structure(project_name: str, outline_json: str) -> str:
        """创建标书项目文件结构"""
        try:
            # 解析大纲JSON
            outline_data = json.loads(outline_json)
            
            # 创建项目目录
            project_dir = f"tender_{project_name.replace(' ', '_')}"
            os.makedirs(project_dir, exist_ok=True)
            os.makedirs(f"{project_dir}/chapters", exist_ok=True)
            os.makedirs(f"{project_dir}/attachments", exist_ok=True)
            os.makedirs(f"{project_dir}/attachments/certificates", exist_ok=True)
            os.makedirs(f"{project_dir}/attachments/cases", exist_ok=True)
            os.makedirs(f"{project_dir}/attachments/images", exist_ok=True)
            os.makedirs(f"{project_dir}/output", exist_ok=True)
            
            # 保存大纲文件
            with open(f"{project_dir}/outline.json", "w", encoding="utf-8") as f:
                json.dump(outline_data, f, ensure_ascii=False, indent=2)
            
            # 创建章节文件
            chapter_files = []
            for chapter in outline_data.get("outline", []):
                chapter_id = chapter.get("chapter_id", "00")
                chapter_title = chapter.get("chapter_title", "未命名章节")
                filename = f"{chapter_id}_{chapter_title.replace('/', '_')}.md"
                filepath = f"{project_dir}/chapters/{filename}"
                
                # 创建章节文件模板
                template_content = f"""# {chapter_title}

## 章节概述
{chapter.get('content_points', [])}

## 主要内容

### 待完善内容
- [ ] 补充具体内容
- [ ] 添加相关图表
- [ ] 完善技术细节

---
*章节状态：未开始*
*预估页数：{chapter.get('estimated_pages', 1)}页*
*权重：{chapter.get('weight', '一般')}*
"""
                
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(template_content)
                
                chapter_files.append(filename)
            
            # 创建项目状态文件
            status_data = {
                "project_name": project_name,
                "created_time": "刚刚创建",
                "total_chapters": len(outline_data.get("outline", [])),
                "completed_chapters": 0,
                "chapter_files": chapter_files,
                "status": "项目已创建，可以开始编写各章节内容"
            }
            
            with open(f"{project_dir}/project_status.json", "w", encoding="utf-8") as f:
                json.dump(status_data, f, ensure_ascii=False, indent=2)
            
            return f"""✅ 标书项目结构创建成功！

📁 项目目录：{project_dir}/
├── outline.json (标书大纲)
├── chapters/ (章节内容，共{len(chapter_files)}个文件)
├── attachments/ (附件目录)
└── output/ (输出目录)

📋 创建的章节文件：
{chr(10).join([f"  - {f}" for f in chapter_files])}

🎯 下一步操作：
1. 使用 tender_content_generator 为各章节生成具体内容
2. 编辑各章节的Markdown文件
3. 使用 merge_tender_chapters 合并所有章节
4. 使用 tender_word_exporter 导出Word文档
"""
        
        except Exception as e:
            return f"❌ 创建项目结构失败：{str(e)}"
    
    @mcp.tool()
    def merge_tender_chapters(project_dir: str) -> str:
        """合并标书章节内容为完整文档"""
        try:
            # 读取大纲文件
            outline_path = f"{project_dir}/outline.json"
            if not os.path.exists(outline_path):
                return f"❌ 找不到大纲文件：{outline_path}"
            
            with open(outline_path, "r", encoding="utf-8") as f:
                outline_data = json.load(f)
            
            # 合并章节内容
            merged_content = []
            
            # 添加封面信息
            tender_info = outline_data.get("tender_info", {})
            merged_content.append(f"""# {tender_info.get('project_name', '标书项目')}

**标书类型：** {tender_info.get('tender_type', '未指定')}
**预估页数：** {tender_info.get('total_pages', '待确定')}
**生成时间：** {outline_data.get('created_time', '未知')}

---

""")
            
            # 添加目录
            merged_content.append("# 目录\n\n")
            for i, chapter in enumerate(outline_data.get("outline", []), 1):
                chapter_title = chapter.get("chapter_title", "未命名章节")
                merged_content.append(f"{i}. {chapter_title}\n")
            merged_content.append("\n---\n\n")
            
            # 合并各章节内容
            chapters_dir = f"{project_dir}/chapters"
            completed_chapters = 0
            
            for chapter in outline_data.get("outline", []):
                chapter_id = chapter.get("chapter_id", "00")
                chapter_title = chapter.get("chapter_title", "未命名章节")
                filename = f"{chapter_id}_{chapter_title.replace('/', '_')}.md"
                filepath = f"{chapters_dir}/{filename}"
                
                if os.path.exists(filepath):
                    with open(filepath, "r", encoding="utf-8") as f:
                        chapter_content = f.read()
                    
                    # 检查章节是否已完成
                    if "*章节状态：已完成*" in chapter_content or len(chapter_content.strip()) > 200:
                        completed_chapters += 1
                    
                    merged_content.append(chapter_content)
                    merged_content.append("\n\n---\n\n")
                else:
                    merged_content.append(f"# {chapter_title}\n\n⚠️ 章节内容缺失\n\n---\n\n")
            
            # 保存合并后的文档
            output_path = f"{project_dir}/output/tender_full.md"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("".join(merged_content))
            
            # 更新项目状态
            status_path = f"{project_dir}/project_status.json"
            if os.path.exists(status_path):
                with open(status_path, "r", encoding="utf-8") as f:
                    status_data = json.load(f)
                
                status_data["completed_chapters"] = completed_chapters
                status_data["last_merged"] = "刚刚合并"
                status_data["merged_file"] = output_path
                
                with open(status_path, "w", encoding="utf-8") as f:
                    json.dump(status_data, f, ensure_ascii=False, indent=2)
            
            return f"""✅ 标书章节合并成功！

📄 合并文件：{output_path}
📊 完成进度：{completed_chapters}/{len(outline_data.get('outline', []))} 章节
📝 文档字数：约 {len(''.join(merged_content))} 字符

🎯 下一步操作：
1. 检查合并后的文档内容
2. 使用 tender_word_exporter 导出Word格式
3. 进行最终的格式调整和完善
"""
        
        except Exception as e:
            return f"❌ 合并章节失败：{str(e)}"
    
    @mcp.tool()
    def export_tender_to_word(project_dir: str, template_style: str = "standard") -> str:
        """将标书导出为Word文档格式"""
        try:
            # 检查合并文件是否存在
            merged_file = f"{project_dir}/output/tender_full.md"
            if not os.path.exists(merged_file):
                return f"❌ 找不到合并文件，请先运行 merge_tender_chapters"
            
            # 读取合并内容
            with open(merged_file, "r", encoding="utf-8") as f:
                markdown_content = f.read()
            
            # 生成Word转换脚本
            word_script = f"""
# Word文档转换脚本
# 需要安装：pip install python-docx markdown

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_to_word():
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 添加标题样式
    styles = doc.styles
    
    # 处理Markdown内容
    lines = '''markdown_content'''.split('\\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # 一级标题
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # 二级标题
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # 三级标题
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            # 粗体段落
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- '):
            # 列表项
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line == '---':
            # 分页符
            doc.add_page_break()
        else:
            # 普通段落
            if line:
                doc.add_paragraph(line)
    
    # 保存文档
    doc.save('{project_dir}/output/tender_full.docx')
    print("Word文档生成成功！")

if __name__ == "__main__":
    convert_to_word()
"""
            
            # 保存转换脚本
            script_path = f"{project_dir}/output/convert_to_word.py"
            with open(script_path, "w", encoding="utf-8") as f:
                f.write(word_script.replace("'''markdown_content'''", repr(markdown_content)))
            
            return f"""✅ Word导出脚本已生成！

📄 转换脚本：{script_path}
📝 源文件：{merged_file}
🎯 目标文件：{project_dir}/output/tender_full.docx

🔧 执行步骤：
1. 安装依赖：pip install python-docx markdown
2. 运行脚本：python {script_path}
3. 获取Word文档：{project_dir}/output/tender_full.docx

💡 样式模板：{template_style}
- standard: 标准标书格式
- professional: 专业商务格式
- government: 政府采购格式

⚠️ 注意事项：
- 请检查生成的Word文档格式
- 根据需要调整样式和布局
- 添加必要的图片和表格
"""
        
        except Exception as e:
            return f"❌ 导出Word文档失败：{str(e)}"
